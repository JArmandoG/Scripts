from csv import DictReader
from docx.shared import Pt
import docx
import sys

def _argchecker():
    try:
        if len(sys.argv) < 3:
            print("[!] Usage: script.py <input.csv> <output.docx>")
            return 0
        if str(sys.argv[1]).endswith("csv") and str(sys.argv[2]).endswith("docx"):
            print("[+] Continuing...")
            generate_report()
            return 1
        else:
            print("[!] Usage: script.py <input.csv> <output.docx>")
            return 0
        # TO-DO: CHECK IF INPUT CSV FILE EXISTS AND RAISE AN ERROR
    except Exception as e:
        print("[XX] An error ocurred")
        print(" > " + str(e))
        return 0


# Valores para el estilo y archivo .csv para procesar
font_name = 'Helvetica'
font_size = 12
INPUT_FILE_CSV = str(sys.argv[1])
OUTPUT_FILE_DOCX = str(sys.argv[2])

def generate_report():
    count = 1
    doc = docx.Document() # Prepara doc como objeto iterable
    with open(INPUT_FILE_CSV, 'r') as csv_file:
        reader_dump = DictReader(csv_file) # Objeto iterable; todavía no es dict
        for my_dict in reader_dump: # Por cada iteración es un dict

            # Generación de tabla y estilos en cada iteración
            table = doc.add_table(rows=11, cols=2)
            style = doc.styles['Normal']
            font = style.font
            font.name = font_name
            font.size = Pt(font_size)
            
            ''' Inline list comprehension (No utilizado)
            keys = [x for x in my_dict.keys()] # Dump KEYS array
            values = [x for x in my_dict.values()] # Dump VALUES array
            '''
            
            '''
            Utilizar DictReader(csv_file).fieldnames para
            visualizar los headers a extraer
            '''
            table.cell(0,0).text = ""
            table.cell(0,1).text = ""

            table.cell(0,2).text = "Control No."
            table.cell(0,3).text = f"{count}"

            # Columna izquierda: String values
            table.cell(0,4).text = "Region"
            # Columna derecha: CSV Fieldnames
            table.cell(0,5).text = my_dict['REGION']

            table.cell(0,6).text = "Severity"
            table.cell(0,7).text = my_dict['CHECK_SEVERITY']
            
            table.cell(0,8).text = "Result"
            table.cell(0,9).text = my_dict['CHECK_RESULT']

            table.cell(0,10).text = "Resource ID"
            table.cell(0,11).text = my_dict['CHECK_RESOURCE_ID']

            table.cell(0,12).text = "Description"
            table.cell(0,13).text = my_dict['TITLE_TEXT']

            table.cell(0,14).text = "Finding"
            table.cell(0,15).text = my_dict['CHECK_RESULT_EXTENDED']

            table.cell(0,16).text = "Associated Risk"
            table.cell(0,17).text = my_dict['CHECK_RISK']

            table.cell(0,18).text = "Remediation"
            table.cell(0,19).text = my_dict['CHECK_REMEDIATION']

            table.cell(0,20).text = "AWS Documentation URL"
            table.cell(0,21).text = my_dict['CHECK_DOC']


            count += 1

        # Salva toda la información guardada en doc.add_table() \ 
        # al objeto docx.Document(), en el documento con el nombre $save_file
        doc.save(OUTPUT_FILE_DOCX)

# generate_report()
_argchecker()
